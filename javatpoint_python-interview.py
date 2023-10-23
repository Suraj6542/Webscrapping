from flask import Flask, render_template, send_file,request, redirect,url_for
from bs4 import BeautifulSoup
import requests
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)


@app.route('/')
def index():
    return render_template('javatpoint.html')

def generation(url,heading):
    response = requests.get(url)
    if response.status_code == 200:
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Find and print specific content (extracting numbered questions)
            s=soup.find('div',id='city')
            headings = s.find_all(['h3'])
            doc = docx.Document()
            h=doc.add_heading(heading, level=1)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
            for heading in headings:
                heading_text = heading.text.strip()
                doc.add_paragraph(heading_text)
            e=doc.add_heading('THE END')
            e.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            filename='interview_questions.docx'
            doc.save(filename)
            return filename

@app.route('/submit/submit',methods=['POST'])
@app.route('/submit', methods=['POST'])
def submit_form():
    if request.method == 'POST':
        url=request.form.get('URL')
        heading=request.form.get('Heading')
        filename = generation(url,heading)
        if filename:
             message=True
        else:
             message=False
        return render_template('javatpoint.html',message=message)
    else:
        return redirect(url_for('index'))
        
    
   
@app.route('/submit/download')
def download():
    filename="C:\\Users\\hp\\OneDrive\\Desktop\\python)\\app\\project\\interview_questions.docx"
    if filename:
        return send_file(filename, as_attachment=True)
    else:
        return redirect(url_for('index'))




    



if __name__ == '__main__':
    app.run(debug=True)
